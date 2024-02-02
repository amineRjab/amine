from PIL import Image  
from PIL import ImageDraw
from PIL import ImageFont
import os
from art import *
from datetime import date
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import random
from docx import Document
from docx.shared import Inches
import os
import time
from art import tprint
import msvcrt
import json

now=datetime.now()
today = date.today()
current_time = now.strftime("%H:%M:%S")
def clear_console():
    
    # Clear the console screen
    os.system('cls' if os.name == 'nt' else 'clear')

def animate_text(text, font='block', speed=1):
    try:
        while True:
            # Clear the console and print the animated text
            clear_console()
            tprint(text, font=font)
            
            # Sleep for a short duration to control the animation speed
            time.sleep(speed)
            
            # Shift the text horizontally for the next frame
            text = ' ' + text[:-1]

            # Check for keyboard input
            if os.name == 'nt':
                if msvcrt.kbhit() and msvcrt.getch() == b'\r':
                    break
            else:
                tty.setcbreak(sys.stdin.fileno())
                if sys.stdin.read(1) == '\r':
                    break
    except KeyboardInterrupt:
        # Handle Ctrl+C to exit the loop gracefully
        pass
    finally:
        # Ensure to restore terminal settings on Unix-like systems
        if os.name != 'nt':
            termios.tcsetattr(sys.stdin, termios.TCSADRAIN, settings)

# Your text to animate
text_to_animate = "Tonz"

# Start the animation
animate_text(text_to_animate, font='random', speed=1)
#302
width = 390
height = 1020


new_data = {}

def insert(date_g,i,start_time,end_time):
    
    if new_data.get(str(date_g)) is not None:
           
            new_data[str(date_g)][str(i)]=str(start_time)+'-'+str(end_time)
           
    else:
            new_data[str(date_g)]={}
            new_data[str(date_g)][str(i)]=str(start_time)+'-'+str(end_time)
            
    print(new_data)

c_start=input("cycle start: ")
c_end=input("cycle end: ")
d_start=input("date start(yyyy-mm-dd hh:mm:s): ")
#6038
time_start = datetime.strptime(d_start, '%d/%m/%Y %H:%M:%S')
counter=0
cycle1=random.randint(8, 11)
cycle2=random.randint(4, 7)
cycle3=random.randint(12, 17)
x=time_start.replace(hour=0, minute=0, second=0,microsecond=0)
x=x+ timedelta(hours=24)
# create a new document

for index,i in enumerate(range(int(c_start),int(c_end)+1)):
    time_start=time_start+ relativedelta(minutes=random.uniform(4, 10))

    
        
        # print('one day')
    if cycle1 == counter:
        img  = Image.new( mode = "RGB", size = (width, height), color = ("white") )
        font = ImageFont.load_default()
        I1 = ImageDraw.Draw(img)
        link=str(os.getcwd())+'\erlinemailoutline.ttf'
        myFont = ImageFont.truetype(link, 40)
        I1.text((35, 10), "NEWSTER (ITALY)", font=myFont, fill =("black"))
        I1.text((5, 50), "NEWSTER 50 S.nr 161", font=myFont, fill =("black"))
        I1.text((62, 90), "NEWSTER", font=myFont, fill =("black"))
        I1.text((87, 130), "ITALY", font=myFont, fill =("black"))
        I1.text((5, 170), "START CYCLE N."+str(i), font=myFont, fill =("black"))
        I1.text((5, 210), "DATE  "+str(time_start.strftime('%d/%m/%y')), font=myFont, fill =("black"))
        I1.text((5, 250), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))
        start_time=str(time_start.time().replace(microsecond=0))
        time_start=time_start+ relativedelta(minutes=random.uniform(0.4, 0.7))
        end_date=str(time_start.strftime('%d/%m/%y'))
        I1.text((5, 340 ), "TEMPERATURE:   60 C", font=myFont, fill =("black"))
        I1.text((5, 380), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(7, 9))
        I1.text((5, 420), "TEMPERATURE:   90 C", font=myFont, fill =("black"))
        I1.text((5, 460), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(33, 37))
        I1.text((5, 500), "TEMPERATURE:  135 C", font=myFont, fill =("black"))
        I1.text((5, 540), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1, 2.5))
        I1.text((5, 580), "TEMPERATURE:  150 C", font=myFont, fill =("black"))
        I1.text((5, 620), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1.8, 3))
        I1.text((5, 660), "TEMPERATURE:  135 C", font=myFont, fill =("black"))
        I1.text((5, 700), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1, 1.5))
        I1.text((5, 740), "TEMPERATURE:   95 C", font=myFont, fill =("black"))
        I1.text((5, 780), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(2, 3.8))
        I1.text((5, 870), "CYCLE ENDED N"+str(i), font=myFont, fill =("black"))
        I1.text((5, 910), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))
        I1.text((5, 950), "MATERIAL STERILIZED", font=myFont, fill =("black"))
        I1.text((40, 990), "-----------------", font=myFont, fill =("black"))

        im1 = img.save(str(i)+".jpg")
        cycle1=random.randint(8, 11)
        counter=11
        insert(str(end_date), i, start_time, str(time_start.time().replace(microsecond=0)))
    elif cycle2 == counter:

        img  = Image.new( mode = "RGB", size = (width, height), color = ("white") )
        font = ImageFont.load_default()
        I1 = ImageDraw.Draw(img)
        link=str(os.getcwd())+'\erlinemailoutline.ttf'
        myFont = ImageFont.truetype(link, 40)
        I1.text((35, 10), "NEWSTER (ITALY)", font=myFont, fill =("black"))
        I1.text((5, 50), "NEWSTER 50 S.nr 161", font=myFont, fill =("black"))
        I1.text((62, 90), "NEWSTER", font=myFont, fill =("black"))
        I1.text((87, 130), "ITALY", font=myFont, fill =("black"))
        I1.text((5, 170), "START CYCLE N."+str(i), font=myFont, fill =("black"))
        I1.text((5, 210), "DATE  "+str(time_start.strftime('%d/%m/%y')), font=myFont, fill =("black"))
        I1.text((5, 250), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))
        start_time=str(time_start.time().replace(microsecond=0))
        time_start=time_start+ relativedelta(minutes=random.uniform(0.4, 2))
        end_date=str(time_start.strftime('%d/%m/%y'))
        I1.text((5, 340 ), "TEMPERATURE:   60 C", font=myFont, fill =("black"))
        I1.text((5, 380), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(4, 10))
        I1.text((5, 420), "TEMPERATURE:   90 C", font=myFont, fill =("black"))
        I1.text((5, 460), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(10, 15))
        I1.text((5, 500), "TEMPERATURE:  135 C", font=myFont, fill =("black"))
        I1.text((5, 540), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1, 5))
        I1.text((5, 580), "TEMPERATURE:  150 C", font=myFont, fill =("black"))
        I1.text((5, 620), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1.8, 3))
        I1.text((5, 660), "TEMPERATURE:  135 C", font=myFont, fill =("black"))
        I1.text((5, 700), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1, 2.5))
        I1.text((5, 740), "TEMPERATURE:   95 C", font=myFont, fill =("black"))
        I1.text((5, 780), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(2, 3.8))
        I1.text((5, 870), "CYCLE ENDED N"+str(i), font=myFont, fill =("black"))
        I1.text((5, 910), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))
        I1.text((5, 950), "MATERIAL STERILIZED", font=myFont, fill =("black"))
        I1.text((40, 990), "-----------------", font=myFont, fill =("black"))

        im1 = img.save(str(i)+".jpg")
        cycle2=random.randint(4, 7)	
        counter=8
        insert(str(end_date), i, start_time, str(time_start.time().replace(microsecond=0)))
    elif cycle3 == counter:
        # print(time_start)
        time_start=time_start+ relativedelta(hours=random.uniform(1, 4))
        cycle3=random.randint(12, 17)
        # print(str(i))
        counter=0
    else:
        img  = Image.new( mode = "RGB", size = (width, height), color = ("white") )
        font = ImageFont.load_default()
        I1 = ImageDraw.Draw(img)
        link=str(os.getcwd())+'\erlinemailoutline.ttf'
        myFont = ImageFont.truetype(link, 40)
        I1.text((35, 10), "NEWSTER (ITALY)", font=myFont, fill =("black"))
        I1.text((5, 50), "NEWSTER 50 S.nr 161", font=myFont, fill =("black"))
        I1.text((62, 90), "NEWSTER", font=myFont, fill =("black"))
        I1.text((87, 130), "ITALY", font=myFont, fill =("black"))
        I1.text((5, 170), "START CYCLE N."+str(i), font=myFont, fill =("black"))
        I1.text((5, 210), "DATE  "+str(time_start.strftime('%d/%m/%y')), font=myFont, fill =("black"))
        I1.text((5, 250), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))
        start_time=str(time_start.time().replace(microsecond=0))
        time_start=time_start+ relativedelta(minutes=random.uniform(0.4, 2))
        end_date=str(time_start.strftime('%d/%m/%y'))
        I1.text((5, 340 ), "TEMPERATURE:   60 C", font=myFont, fill =("black"))
        I1.text((5, 380), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(5, 10))
        I1.text((5, 420), "TEMPERATURE:   90 C", font=myFont, fill =("black"))
        I1.text((5, 460), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(15, 20))
        I1.text((5, 500), "TEMPERATURE:  135 C", font=myFont, fill =("black"))
        I1.text((5, 540), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(2, 5))
        I1.text((5, 580), "TEMPERATURE:  150 C", font=myFont, fill =("black"))
        I1.text((5, 620), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1.8, 3))
        I1.text((5, 660), "TEMPERATURE:  135 C", font=myFont, fill =("black"))
        I1.text((5, 700), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(1, 2.5))
        I1.text((5, 740), "TEMPERATURE:   95 C", font=myFont, fill =("black"))
        I1.text((5, 780), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))

        time_start=time_start+ relativedelta(minutes=random.uniform(2, 3.8))
        
        I1.text((5, 870), "CYCLE ENDED N"+str(i), font=myFont, fill =("black"))
        I1.text((5, 910), "TIME: "+str(time_start.time().replace(microsecond=0)), font=myFont, fill =("black"))
        I1.text((5, 950), "MATERIAL STERILIZED", font=myFont, fill =("black"))
        I1.text((40, 990), "-----------------", font=myFont, fill =("black"))

        im1 = img.save(str(i)+".jpg")
        counter=counter+1
        insert(str(end_date), i, start_time, str(time_start.time().replace(microsecond=0)))



img.show()

# Loop through the outer dictionary
for outer_key, inner_dict in new_data.items():
        # print(f"Outer Key: {outer_key}")
        doc = Document()

                # add a heading to the document
        doc.add_heading(outer_key, level=0)
                # add a table to the document with 3 columns
        table = doc.add_table(rows=1,cols=4)

                # add headers to the table
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'N cycle'
        hdr_cells[1].text = 'H.Debit'
        hdr_cells[2].text = 'H.Fin'
        hdr_cells[3].text = 'Remarque'
    # Loop through the inner dictionary
        for inner_key, value in inner_dict.items():
            
        
      
            folder_name = 'doc'
            if not os.path.exists(folder_name):
                os.mkdir(folder_name)
            # add data to the table
            row_cells = table.add_row().cells
            row_cells[0].text = str(inner_key)
            p=value.split('-')
            row_cells[1].text = p[0]
            row_cells[2].text = p[1]
            print(f"Inner Key: {inner_key}, Value: {value}")
            
        row_cells = table.add_row().cells
        row_cells[0].text = str(inner_key)
        row_cells[1].text = '-'
        row_cells[2].text = '-'
        file_name=outer_key.replace('/', '-')+'.docx'
            # save the document to the folder
            
        file_path = os.path.join(folder_name, file_name)
            
        doc.save(file_path)
            
# shared_module.py
new_data = {}

def insert(date_g, i, start_time, end_time):
    if new_data.get(str(date_g)) is not None:
        new_data[str(date_g)][str(i)] = str(start_time) + '-' + str(end_time)
    else:
        new_data[str(date_g)] = {}
        new_data[str(date_g)][str(i)] = str(start_time) + '-' + str(end_time)
    print(new_data)
           
